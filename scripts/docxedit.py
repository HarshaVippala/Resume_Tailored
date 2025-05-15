from docx import Document
import re

def replace_string(doc, old_string, new_string):
    """
    Replace all occurrences of old_string with new_string in the document,
    handling cases where the placeholder might be split across multiple runs.
    Preserve styles of the original runs as much as possible.
    """
    def process_paragraph(paragraph):
        # Find all runs that together contain the placeholder
        text = ''.join(run.text for run in paragraph.runs)
        if old_string not in text:
            return False
        
        # Find the start and end run indices for the placeholder
        joined = ''
        start_idx = end_idx = None
        for i, run in enumerate(paragraph.runs):
            if start_idx is None and old_string.startswith(run.text):
                joined = run.text
                start_idx = i
                if joined == old_string:
                    end_idx = i
                    break
            elif start_idx is not None:
                joined += run.text
                if joined == old_string:
                    end_idx = i
                    break
        # If not found as split, fallback to simple replace in one run
        if start_idx is None or end_idx is None:
            for run in paragraph.runs:
                if old_string in run.text:
                    run.text = run.text.replace(old_string, new_string)
                    return True
            return False
        # Merge runs and replace
        first_run = paragraph.runs[start_idx]
        # Concatenate text before, replace, and after
        before = ''.join(run.text for run in paragraph.runs[:start_idx])
        after = ''.join(run.text for run in paragraph.runs[end_idx+1:])
        new_full = before + new_string + after
        # Remove all runs
        for _ in range(len(paragraph.runs)):
            paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
        # Add new run with merged text, copy style from first_run
        new_run = paragraph.add_run(new_full)
        new_run.bold = first_run.bold
        new_run.italic = first_run.italic
        new_run.underline = first_run.underline
        new_run.font.size = first_run.font.size
        new_run.font.name = first_run.font.name
        new_run.style = first_run.style
        return True
    # Process all paragraphs in the document
    for para in doc.paragraphs:
        process_paragraph(para)
    # Also handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

def extract_placeholders(doc):
    # Find all unique placeholders in the document (e.g., <PLACEHOLDER>)
    placeholders = set()
    pattern = re.compile(r'<[A-Z0-9_&]+>')
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        for match in pattern.findall(para.text):
            placeholders.add(match)
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for match in pattern.findall(para.text):
                        placeholders.add(match)
    
    return list(placeholders)

# Check if placeholders are split across runs and merge them
def preprocess_document(doc):
    """
    Preprocess document to identify and merge runs that might contain partial placeholders.
    This helps when placeholders are split across runs due to formatting.
    """
    pattern = re.compile(r'<[A-Z0-9_&]+>')
    
    # Process paragraphs
    for para in doc.paragraphs:
        # If there's a potential placeholder marker '<' or partial match
        if any('<' in run.text for run in para.runs) and len(para.runs) > 1:
            # Get the paragraph text and look for placeholders
            text = para.text
            placeholders_in_para = pattern.findall(text)
            
            # If there are placeholders but they're not in individual runs, we need to merge
            if placeholders_in_para and not any(ph in run.text for ph in placeholders_in_para for run in para.runs):
                # Clear and recreate runs
                for i in range(len(para.runs)):
                    para.runs[0]._element.getparent().remove(para.runs[0]._element)
                para.add_run(text)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if any('<' in run.text for run in para.runs) and len(para.runs) > 1:
                        text = para.text
                        placeholders_in_para = pattern.findall(text)
                        
                        if placeholders_in_para and not any(ph in run.text for ph in placeholders_in_para for run in para.runs):
                            for i in range(len(para.runs)):
                                para.runs[0]._element.getparent().remove(para.runs[0]._element)
                            para.add_run(text)
    
    return doc 