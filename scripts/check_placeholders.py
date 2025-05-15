from docx import Document
import re
import sys

def extract_placeholders(doc):
    # Find all unique placeholders in the document (e.g., <PLACEHOLDER>)
    placeholders = set()
    pattern = re.compile(r'<[A-Z0-9_&]+>')
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        for match in pattern.findall(para.text):
            placeholders.add(match)
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for match in pattern.findall(para.text):
                        placeholders.add(match)
    
    return list(placeholders)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python check_placeholders.py <docx_file>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    try:
        doc = Document(file_path)
        placeholders = extract_placeholders(doc)
        print(f"\nFound {len(placeholders)} placeholders in {file_path}:")
        for ph in sorted(placeholders):
            print(f"  {ph}")
    except Exception as e:
        print(f"Error: {e}") 