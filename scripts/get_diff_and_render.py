import json
import sys
from openai import OpenAI
from make_resume import patch_docx
from dotenv import load_dotenv
import os
from docx2pdf import convert
import re
import argparse
from docxedit import extract_placeholders

load_dotenv()

def extract_json_from_markdown(text):
    match = re.search(r"```(?:json)?\s*([\s\S]+?)\s*```", text, re.IGNORECASE)
    if match:
        return match.group(1)
    return text

def get_diff_from_gpt(jd_path, template_path, base_path, api_key=None):
    client = OpenAI(api_key=api_key)
    job_desc = open(jd_path).read()
    
    from docx import Document
    # Extract placeholders from the template
    template_doc = Document(template_path)
    placeholders = extract_placeholders(template_doc)
    # Build a JSON skeleton and code block
    json_skeleton = '{\n' + ',\n'.join([f'  "{ph}": ""' for ph in placeholders]) + '\n}'
    json_template = "```json\n" + json_skeleton + "\n```"
    placeholder_keys = ', '.join(placeholders)

    # Extract resume text from the base resume
    base_doc = Document(base_path)
    resume_text = '\n'.join([para.text for para in base_doc.paragraphs])

    prompt = (
        "You are CareerForgeAI, an elite career strategist and resume optimization specialist with 15+ years of executive recruitment experience across Fortune 500 companies and specialized in applicant tracking systems (ATS) algorithms."
        "Modern hiring processes rely heavily on automated screening and psychological triggers that determine which candidates advance. 85 percent of resumes are rejected before human eyes ever see them. Standard resume advice fails to address the technical and psychological aspects of successful applications."
        "Your task is to conduct a deep analysis of the provided job description and the candidate's base resume (already stored in memory) to identify technical and psychological gaps. Produce an ATS-optimized resume with properly weighted keywords in STAR format, tailored to the job description, while ensuring the content bypasses AI detection by mimicking human writing. Populate a JSON object with the updated text content for the summary, skills, and work experience sections. The output MUST be a single, flat JSON object."
        #### CRITICAL INSTRUCTIONS FOR JSON KEY FORMATTING:
        "1. You MUST use the EXACT placeholder keys as provided in the `JSON_TEMPLATE` below."
        "2. Keys are case-sensitive and character-sensitive. They MUST include the angle brackets `<` and `>` and be in ALL CAPS or the exact case as shown in the template (e.g., `<SUMMARY>`, `<JOB1_POINT1>`)."
        "3. DO NOT modify the key names in any way. This means:"
        "- NO converting to lowercase."
        "- NO converting to snake_case or camelCase."
        "- NO removing or changing angle brackets or any other characters."
        "- NO adding, removing, or renaming any keys from the template."
        "4. The entire response MUST be ONLY the JSON object, starting with `{` and ending with `}`. Do not include any text before or after the JSON object, including markdown code fences."
        #### JSON_TEMPLATE (Fill in the empty string values for each key. Preserve keys EXACTLY as shown):
        "{json_template}"
        #### JSON_SKELETON_TO_POPULATE (Fill in the empty string values for each key. Preserve keys EXACTLY as shown. This is the structure your JSON output must follow):
        "{json_skeleton}"
        #### CONTENT GUIDELINES (for the string values in the JSON):
        "- If no specific information is available for a placeholder key, use an empty string `""` as its value."
        "- All values associated with keys MUST be strings. Do NOT use nested JSON objects or JSON arrays as values."
        "- For placeholders representing a list of points (e.g., for job experience bullet points like `<JOB1_POINT1>`, `<JOB1_POINT2>`), each such key should receive content for its corresponding single point. If a single key is intended to hold multiple distinct points, combine them into a single string with each point on a new line (separated by `\n`)."
        "- For SKILLS placeholders (e.g., `<SKILLS_CLOUDDEVOPS>`, `<SKILLS_MONITORINGOBSERVABILITY>`): List specific and discrete software, technologies, tools, libraries, frameworks, and well-defined methodologies (e.g., 'Python', 'React', 'AWS Lambda', 'Docker', 'Git', 'Agile', 'Scrum'). Do NOT list general concepts, practices, or categories (e.g., avoid terms like 'cloud computing', 'data analysis', 'software development', 'application performance monitoring' as standalone items unless they are part of a specific, named methodology or platform you are listing). "
        "Focus on concrete nouns that represent actual skills a person uses."
        "**ATS Optimization:**"
        "- Analyze the job description to identify key skills, technologies, and requirements."
        "- Update the summary section to highlight the most relevant experience and skills matching the job."
        "- Tailor the skills section to reflect the job requirements, prioritizing the most relevant technologies and methodologies."   
        "- Rewrite work experience bullet points to emphasize achievements and technologies relevant to the job, using the STAR format (Situation, Task, Action, Result) and quantifiable metrics where possible."  
        "- Naturally incorporate relevant keywords from the job description into the resume content, ensuring they fit seamlessly into the text."   
        "**Bypassing AI Detection:**"
        "- Use specific, personalized language and avoid generic phrases or buzzwords (e.g., steer clear of 'results-driven' or 'team player' unless uniquely contextualized)."
        "- Incorporate personal anecdotes or details from the candidate's experience to make the resume unique (e.g., specific project names like '7-Eleven Store Analytics Platform' or personal motivations)."
        "- Vary sentence structure and length to mimic natural human writing, mixing short and long sentences for increased burstiness."
        "- Use a mix of technical and non-technical language to sound authentic (e.g., blend 'Node.js' with 'streamlined workflows')."
        "- Occasionally include unique expressions or less common word choices to enhance the human-like quality (e.g., 'I've delved into' instead of 'I have experience in')."
        "- In the summary section, include a sentence about the candidate's approach to software engineering or their professional philosophy (e.g., 'I thrive on crafting scalable solutions that balance performance and maintainability')."
        "- In work experience bullet points, include a brief statement about the impact or significance of the work (e.g., 'This effort slashed transaction times, delighting customers')."
        "- Use a professional tone with a hint of enthusiasm for technology and problem-solving."
        "- Use specific project names, technologies, and achievements from the candidate's existing resume, enhancing them with personalized language and impact statements."
        "- Focus on concise, relevant bullet points."
        "- Work experience bullet points MUST reference the specific project name or context, list the primary technologies/frameworks used, and include a quantifiable outcome or metric (e.g., 'On [ProjectName] using [TechStack], achieved [Result] resulting in [Metric]'). Do not reference projects, teams, or activities from any other employer."
        "- In the SUMMARY section, specify an exact number of years of experience (either 4 or 5 years); do not use the '+' notation."
        "- DO NOT mention the target company's name or systems in work experience bullet points (e.g., avoid phrases like 'skills critical for [Company]'s core systems')."
        #### CHARACTER LIMIT CONSTRAINTS (CRITICAL FOR PROPER FORMATTING):
        "- SUMMARY section: MUST be exactly 369 characters (3 lines * 123 characters per line) including spaces. This ensures the summary fits precisely on 3 lines without overflow."
        "- SKILLS sections: Maximum 7 skills per category, listing most important skills first. Skills should be presented as comma-separated values on a single line. Architecture and design bullet in skills section should not have more than 4 values."
        "- WORK EXPERIENCE bullet points: Each bullet point MUST NOT exceed 246 characters (2 lines * 123 characters per line) including spaces. Aim for a mix of concise and detailed points within this range, with some bullets being more succinct (~180-200 chars) and others more comprehensive (~210-246 chars) as the content requires. Do not use the '+', ';' and '-' in the bullet points."
        #### ADDITIONAL INSTRUCTIONS:
        "- Keep the count of 7-Eleven stores less than 2500 stores for work experience bullet points."
        "- Languages and Frameworks should have a mix of languages and their related frameworks. Do not mention HTML and CSS as a language and framework."
        "- Use varied, natural sentence structures."
        "- Be technically specific and accurate: mention technologies, tools, frameworks, and methodologies relevant to the job description."
        "- Use project-specific details from the base resume (e.g., '7-Eleven Store Analytics Platform', 'Retail Monitoring Dashboard', 'Claims Processing Pipeline' for Liberty Mutual) with their associated technologies and frameworks." 
        "- If creating new projects, align them with company operations (e.g., '7-Eleven Inventory Management System', 'Liberty Mutual Policy Validation Service')."
        "- For each project, specify the actual technologies used (e.g., OpenTelemetry, AWS CloudWatch, Node.js, TypeScript) and include quantifiable performance metrics and business impact."
        "- Select the most relevant roles for the target job description. It's acceptable if 1 or 2 work experience bullets are not directly related to the job description; prioritize showcasing core strengths and impact."
        "- Include concrete, quantifiable achievements with metrics where possible."
        "- Use '%' symbol instead of spelling out 'percent' in all metrics and statistics (e.g., '22%' instead of '22 percent')."
        "- Avoid formulaic or AI-detectable language."
        "- Check grammar and spelling of the output."
        #### INPUTS:
        "- **Base Resume Text:** The candidate's current resume, containing summary, skills, and work experience details (e.g., roles at 7-Eleven, Liberty Mutual)."
        "- **Job Description (provided as input):** The job posting to tailor the resume for."
        #### OUTPUT:
        "A single, flat JSON object populated with tailored text content for the summary, skills, and work experience sections, adhering to all guidelines and constraints."
    )
    
    # Include the base resume text directly in the prompt
    prompt = prompt.replace("- **Base Resume Text:** The candidate's current resume", 
                           f"- **Base Resume Text:**\n{resume_text}\n\n- **Job Description (provided as input):**\n{job_desc}")
    
    # Ensure json_template and json_skeleton are correctly formatted
    prompt = prompt.replace("{json_template}", json_template)
    prompt = prompt.replace("{json_skeleton}", json_skeleton)
    
    response = client.chat.completions.create(
        model="gpt-4.1",
        messages=[{"role": "user", "content": prompt}]
    )
    
    content = response.choices[0].message.content
    try:
        diff_data = json.loads(content)
    except json.JSONDecodeError:
        clean_content = extract_json_from_markdown(content)
        diff_data = json.loads(clean_content)
    # Check for nested/sectioned output and retry if necessary
    if any(isinstance(v, (list, dict)) for v in diff_data.values()):
        print("Detected nested or sectioned output from LLM. Retrying with explicit flat mapping instructions...")
        retry_prompt = (
            "You must return a FLAT JSON mapping where each key matches the placeholder format (e.g., <SUMMARY>, <JOB1_POINT1>, etc.).\n"
            "Do NOT use sections, arrays, or change key names.\n"
            f"JSON template (fill in the values):\n{json_skeleton}\n\n"
            f"Base Resume:\n{resume_text}\n\n"
            f"Job Description:\n{job_desc}\n"
        )
        retry_response = client.chat.completions.create(
            model="gpt-4.1",
            messages=[{"role": "user", "content": retry_prompt}]
        )
        retry_content = retry_response.choices[0].message.content
        try:
            diff_data = json.loads(retry_content)
        except json.JSONDecodeError:
            clean_retry = extract_json_from_markdown(retry_content)
            diff_data = json.loads(clean_retry)
    
    # Verify that all placeholders are included
    missing_placeholders = [p for p in placeholders if p not in diff_data]
    if missing_placeholders:
        missing_json_skeleton = '{\n' + ',\n'.join([f'  "{ph}": ""' for ph in missing_placeholders]) + '\n}'
        missing_list = ', '.join(missing_placeholders)
        print(f"Warning: The following placeholders were not generated: {missing_list}")
        
        # Make another API call to fill missing placeholders
        missing_prompt = (
            f"You previously generated content for a resume but missed the following placeholders: {missing_list}\n\n"
            "Please generate content for ONLY these missing placeholders, considering the job description and base resume. "
            "Return ONLY a valid JSON with these placeholder keys mapped to optimized content strings. DO NOT OMIT ANY KEY.\n\n"
            
            "**CHARACTER LIMIT CONSTRAINTS (CRITICAL FOR PROPER FORMATTING):**\n"
            "- SUMMARY section: Must be between 400-500 characters (including white spaces). Concise yet comprehensive overview of professional background.\n"
            "- SKILLS sections: Maximum 7 skills per category, listing most important skills first. Skills should be presented as comma-separated values on a single line.\n"
            "- WORK EXPERIENCE bullet points: Each bullet must be between 120-240 characters (including white spaces). Include metrics and achievements while maintaining this length constraint.\n\n"
            
            f"JSON template (fill in the values):\n{missing_json_skeleton}\n\n"
            f"Base Resume:\n{resume_text}\n\n"
            f"Job Description:\n{job_desc}\n"
        )
        
        missing_response = client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": content},
                {"role": "user", "content": missing_prompt}
            ]
        )
        
        missing_content = missing_response.choices[0].message.content
        try:
            # Try to parse directly first
            missing_data = json.loads(missing_content)
        except json.JSONDecodeError:
            # If direct parsing fails, try to extract JSON from markdown code blocks
            clean_missing = extract_json_from_markdown(missing_content)
            missing_data = json.loads(clean_missing)
        
        # Merge the two sets of data
        diff_data.update(missing_data)
    
    return json.dumps(diff_data, indent=2)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate a diff for resume tailoring")
    parser.add_argument("--jd", required=True, help="Path to job description file")
    parser.add_argument("--template", required=True, help="Path to template resume")
    parser.add_argument("--base", required=True, help="Path to base resume")
    parser.add_argument("--diff", required=True, help="Path to save diff JSON")
    parser.add_argument("--output", help="Path to save output resume (if specified)")
    
    args = parser.parse_args()
    
    api_key = os.getenv("OPENAI_API_KEY")
    diff_data = get_diff_from_gpt(args.jd, args.template, args.base, api_key)
    
    # Save diff to a file
    with open(args.diff, "w") as f:
        f.write(diff_data)
    
    # If output path is provided, generate the resume
    if args.output:
        template_path = os.path.join('data', 'master_resume.dotx')
        patch_docx(
            template_path=template_path,
            base_path=args.base,
            diff_path=args.diff,
            output_path=args.output
        )
        
        # Set permissions on the generated DOCX file to user read/write (chmod 644)
        try:
            os.chmod(args.output, 0o644)
        except Exception as e:
            print(f"Warning: Failed to set permissions on DOCX file: {e}")
        
        # Optionally convert to PDF
        try:
            pdf_path = os.path.splitext(args.output)[0] + ".pdf"
            convert(args.output, pdf_path)
            # Set permissions on the generated PDF file to user read/write (chmod 644)
            try:
                os.chmod(pdf_path, 0o644)
            except Exception as e:
                print(f"Warning: Failed to set permissions on PDF file: {e}")
            print(f"Generated PDF: {pdf_path}")
        except Exception as e:
            print(f"Warning: Failed to convert to PDF with docx2pdf: {e}")
            # Fallback: Try AppleScript via osascript for PDF conversion
            import subprocess
            applescript = f'''
            tell application "Microsoft Word"
                open POSIX file "{os.path.abspath(args.output)}"
                set theDoc to active document
                set pdfPath to "{os.path.abspath(pdf_path)}"
                save as theDoc file format format PDF file name pdfPath
                close theDoc saving no
            end tell
            '''
            try:
                result = subprocess.run([
                    "osascript", "-e", applescript
                ], capture_output=True, text=True)
                if result.returncode == 0:
                    print(f"[Fallback] PDF generated via AppleScript: {pdf_path}")
                    try:
                        os.chmod(pdf_path, 0o644)
                    except Exception as e:
                        print(f"Warning: Failed to set permissions on fallback PDF file: {e}")
                else:
                    print(f"[Fallback] AppleScript PDF conversion failed: {result.stderr}")
            except Exception as ase:
                print(f"[Fallback] Exception during AppleScript PDF conversion: {ase}")
            
    print(f"Diff saved to {args.diff}")
    if args.output:
        print(f"Tailored resume saved to {args.output}")
