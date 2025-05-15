import json
import sys
import os
from docx import Document
from docxedit import replace_string, extract_placeholders, preprocess_document


def extract_base_mapping(base_path):
    doc = Document(base_path)
    mapping = {}
    placeholders = extract_placeholders(doc)
    for ph in placeholders:
        # Use the text in the doc as the value for the placeholder
        for para in doc.paragraphs:
            if ph in para.text:
                mapping[ph] = para.text.strip()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if ph in para.text:
                            mapping[ph] = para.text.strip()
    return mapping

def bold_skill_labels(doc):
    """
    Re‑apply bold formatting to skill category labels only.
    
    This function identifies skill headings by looking for paragraphs that:
    1. Contain "SKILLS_" placeholder or 
    2. Begin with common skill category prefixes followed by a colon
    
    Only the text before (and including) the colon is bolded, leaving
    the skill values in normal formatting.
    """
    # Common skill category prefixes to identify skill sections
    skill_prefixes = [
        "Languages & Frameworks", 
        "Cloud & DevOps", 
        "APIs & Integration",
        "Architecture & Design",
        "Databases & Storage",
        "Monitoring & Observability",
        "Testing & CI/CD",
        # Add other skill categories that might be in your resume format
    ]
    
    for para in doc.paragraphs:
        # Skip paragraphs without colons
        if ':' not in para.text:
            continue
            
        # Check if this is a skills paragraph by looking for:
        # 1. SKILLS_ placeholder pattern
        # 2. Common skill category prefix
        is_skills_para = False
        if any(placeholder in para.text for placeholder in ["<SKILLS_", "<skills_"]):
            is_skills_para = True
        else:
            text_before_colon = para.text.split(':', 1)[0].strip()
            if any(prefix in text_before_colon for prefix in skill_prefixes):
                is_skills_para = True
                
        # Skip non-skills paragraphs
        if not is_skills_para:
            continue
            
        # Bold the text before and including the colon
        colon_seen = False
        for run in para.runs:
            if colon_seen:
                break
            if ':' in run.text:
                colon_index = run.text.find(':')
                if colon_index < len(run.text) - 1:
                    # If colon is not at the end of the run, we need to split the run
                    # This preserves bold formatting for text before/including colon only
                    part1 = run.text[:colon_index+1]
                    part2 = run.text[colon_index+1:]
                    run.text = part1
                    run.bold = True
                    
                    # Create a new run for the text after the colon
                    new_run = para.add_run(part2)
                    new_run.bold = False
                    # Copy other formatting (except bold)
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                else:
                    # Colon is at the end of the run, just bold the whole run
                    run.bold = True
                colon_seen = True
            else:
                run.bold = True

def patch_docx(template_path, diff_json, base_path, out_path):
    # Handle .dotx files by using a regular docx file instead
    if template_path.lower().endswith('.dotx'):
        print(f"Warning: Template file {template_path} is a .dotx file which may not be directly supported.")
        print("Using base resume as template and applying placeholder replacements.")
        doc = Document(base_path)
    else:
        doc = Document(template_path)
    
    # Preprocess document to handle split placeholders
    doc = preprocess_document(doc)
    
    base_mapping = extract_base_mapping(base_path)
    placeholders = extract_placeholders(doc)
    
    # Track which placeholders were replaced
    replaced = set()
    
    for ph in placeholders:
        if ph in diff_json:
            replace_string(doc, ph, diff_json[ph])
            replaced.add(ph)
        elif ph in base_mapping:
            replace_string(doc, ph, base_mapping[ph])
            replaced.add(ph)
        else:
            print(f"Warning: Placeholder {ph} not found in diff or base resume, leaving as is.")

    bold_skill_labels(doc)  # ensure skill headings stay bold

    # Log replacement details
    print(f"Found {len(placeholders)} placeholders in template:")
    for ph in placeholders:
        status = "✓ Replaced" if ph in replaced else "⚠ Not replaced"
        print(f"  {ph}: {status}")
    
    doc.save(out_path)
    print(f"\nReplaced {len(replaced)} placeholders out of {len(placeholders)} found.")

    # Post-processing check for unreplaced placeholders
    remaining_placeholders = extract_placeholders(doc)
    if remaining_placeholders:
        print(f"\nERROR: The following placeholders were NOT replaced:")
        for ph in remaining_placeholders:
            print(f"  {ph}")
        print("\nPlease check your diff and base resume for missing keys.")
        sys.exit(1)
    else:
        print("\nAll placeholders successfully replaced.")

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Patch a resume template with diff data")
    parser.add_argument("--template", required=True, help="Path to the template file (.docx or .dotx)")
    parser.add_argument("--diff", required=True, help="Path to the diff JSON file")
    parser.add_argument("--base", required=True, help="Path to the base resume file")
    parser.add_argument("--output", required=True, help="Path to save the output file")
    
    args = parser.parse_args()
    
    with open(args.diff) as f:
        diff = json.load(f)
    
    patch_docx(args.template, diff, args.base, args.output)

# commit: update patch_docx to handle .dotx files, fix placeholder replacement across runs, and improve debugging output
